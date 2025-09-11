"""
Bulletin router for generating bulletin DOCX files
"""
from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse, Response
from typing import List, Dict, Any
import os
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_COLOR_INDEX

router = APIRouter()

@router.post("/generate-bulletin")
async def generate_bulletin(data: Dict[str, Any]):
    """Generate bulletin DOCX from bulletin data"""
    return await export_docx(data)

@router.post("/generate-docx")
async def generate_docx(data: Dict[str, Any]):
    """Legacy endpoint - redirects to generate-bulletin"""
    return await export_docx(data)

def create_bulletin_styles(doc):
    """Create custom styles for bulletin document."""
    styles = doc.styles
    
    # Header style
    header_style = styles.add_style('HeaderStyle', 1)  # 1 = paragraph style
    header_style.font.name = 'Trade Gothic Next Cond'
    header_style.font.size = Pt(10)
    header_style.font.italic = True
    header_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_style.paragraph_format.space_after = Pt(0)
    header_style.paragraph_format.space_before = Pt(0)
    
    # Body style for table content
    body_style = styles.add_style('BodyStyle', 1)  # 1 = paragraph style
    body_style.font.name = 'Trade Gothic Next Cond'
    body_style.font.size = Pt(12)
    body_style.paragraph_format.space_after = Pt(2.16)
    body_style.paragraph_format.space_before = Pt(0)
    
    # Small empty lines style
    small_style = styles.add_style('SmallEmpty', 1)  # 1 = paragraph style
    small_style.font.name = 'Trade Gothic Next Cond'
    small_style.font.size = Pt(6)
    small_style.paragraph_format.space_after = Pt(0)
    small_style.paragraph_format.space_before = Pt(0)
    
    return header_style, body_style, small_style


def add_bulletin_item(doc, item, center_tab=3.1, right_tab=6.2):
    """Add a single bulletin item to the document."""
    # Extract the required fields
    standing = item.get('Standing', False)
    text = item.get('Text', '').strip()  # Original text from HTML
    name = item.get('Name', '').strip()  # Template name if matched
    description = item.get('Description', '').strip()
    personnel = item.get('Personnel', '').strip()
    was_matched = item.get('wasMatched', False)
    
    # For matched items, use the template name in uppercase
    # For unmatched items, use the original text
    if was_matched and name:
        display_text = name.upper()
        
        # Special cases for hymns based on template name
        if name.lower() == "opening hymn":
            display_text = "HYMN"  # First hymn is just "HYMN"
        elif name.lower() == "closing hymn":
            display_text = "CLOSING HYMN"
        elif name.lower() == "hymn":
            display_text = "HYMN"  # Middle hymn
    else:
        # Use original text for unmatched items
        display_text = text
        if ':' in text and description:
            # Remove colon part if we have a description
            display_text = text.split(':')[0].strip()
    
    # Special handling for specific items
    if "BLOWING OF THE SHELL" in display_text.upper():
        display_text = "BLOWING OF THE SHELL (pu in Hawaiian, kele'a in Tongan)"
        description = ""  # Clear the "(PU)" part from description
    elif "FAITH" in text.upper() and "HOPE" in text.upper():
        display_text = "FAITH IN ACTION STEPS AND INTRODUCTION OF OFFERING"
    
    # Fix "SHARING OF" duplicate in description
    if description and "SHARING OF" in description and "SHARING OF" in text:
        description = ""
    
    # Add standing indicator
    if standing:
        display_text = "*" + display_text
    
    # Create a new paragraph with body style
    paragraph = doc.add_paragraph(style='BodyStyle')
    
    # Clear any default tab stops
    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stops.clear_all()
    
    # Check if description is empty or None
    if not description or description.lower() == 'none':
        # Single tab stop for right alignment when no description
        tab_stops.add_tab_stop(Inches(right_tab), alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.DOTS)
        
        # Add text with single tab
        run = paragraph.add_run(display_text + '\t' + personnel)
    else:
        # Two tab stops when description is present
        tab_stops.add_tab_stop(Inches(center_tab), alignment=WD_TAB_ALIGNMENT.CENTER, leader=WD_TAB_LEADER.DOTS)
        tab_stops.add_tab_stop(Inches(right_tab), alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.DOTS)
        
        # Add text with two tabs
        run = paragraph.add_run(display_text + '\t' + description + '\t' + personnel)
    
    # Highlight the line if personnel is missing (yellow background)
    if not personnel and was_matched:
        # Check if this is an item that should have personnel
        items_needing_personnel = [
            'opening prayer', 'scripture', 'sermon', 'message for all generations',
            'sharing of joys', 'prayer of dedication', 'benediction', 'call to worship'
        ]
        if any(item in name.lower() for item in items_needing_personnel):
            # Apply yellow highlight to the entire paragraph
            for run in paragraph.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        
    # Add Children release line for Message for All Generations
    if "MESSAGE FOR ALL GENERATIONS" in display_text.upper():
        # Create a new paragraph with body style
        paragraph = doc.add_paragraph(style='BodyStyle')
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run("Children are released to Sunday School. The teacher today is ????????")
        run.font.italic = True


@router.post("/export/docx")
async def export_docx(data: Dict[str, Any]):
    """Export bulletin data as DOCX with proper formatting"""
    try:
        # Handle both 'bulletinData' and 'data' keys for compatibility
        bulletin_data = data.get('bulletinData') or data.get('data', [])
        service_date = data.get('serviceDate', datetime.now().strftime('%Y-%m-%d'))
        filename = data.get('filename', f'bulletin_{service_date}.docx')
        
        # Create a new Document
        doc = Document()
        
        # Create custom styles
        create_bulletin_styles(doc)
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.1)
            section.right_margin = Inches(1.1)
        
        # Add header lines
        for text in [
            " ", " ",  # Empty lines
            "* Please rise as you are able in body and/or spirit.",
            " "   # Empty line
        ]:
            doc.add_paragraph(text, style='HeaderStyle')
        
        # Process bulletin items
        for item in bulletin_data:
            add_bulletin_item(doc, item)
        
        # Add footer - 15 empty lines
        for i in range(15):
            doc.add_paragraph(" ", style='SmallEmpty')
        
        # Save to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name
        
        # Return the file
        return FileResponse(
            path=tmp_path,
            filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

